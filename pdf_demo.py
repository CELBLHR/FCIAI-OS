# -*- coding: utf-8 -*-
import os
import sys
import zipfile
import logging
import traceback
from pathlib import Path

import requests
import certifi
import pypandoc
from docx import Document
from docx.enum.text import WD_BREAK


logger = logging.getLogger("pdf_demo")
handler = logging.StreamHandler(sys.stdout)
handler.setFormatter(logging.Formatter("[%(levelname)s] %(message)s"))
logger.addHandler(handler)
logger.setLevel(logging.INFO)


def safe_mkdir(p: str | Path):
    Path(p).mkdir(parents=True, exist_ok=True)


def download_or_copy_zip(zip_url: str, zip_path: str):
    if zip_url.startswith("file://"):
        src = zip_url[7:]
        if not os.path.exists(src):
            raise FileNotFoundError(f"源文件不存在: {src}")
        import shutil
        shutil.copy2(src, zip_path)
        return

    resp = requests.get(zip_url, timeout=300, verify=certifi.where())
    if resp.status_code != 200:
        raise RuntimeError(f"下载ZIP失败，状态码: {resp.status_code}，响应: {resp.text[:300]}")
    with open(zip_path, "wb") as f:
        f.write(resp.content)


def extract_zip(zip_path: str, out_dir: str) -> list[str]:
    with zipfile.ZipFile(zip_path, "r") as zf:
        file_list = zf.namelist()
        zf.extractall(out_dir)
    return file_list


def find_md_or_txt(root_dir: str, task_id: str | None = None) -> str | None:
    candidates: list[str] = []
    for r, _, files in os.walk(root_dir):
        for fn in files:
            rel = os.path.relpath(os.path.join(r, fn), root_dir)
            candidates.append(rel)

    if task_id:
        for rel in candidates:
            if rel.lower().endswith(".md") and task_id in rel:
                return os.path.join(root_dir, rel)

    for rel in candidates:
        if rel.lower().endswith(".md"):
            return os.path.join(root_dir, rel)

    for rel in candidates:
        if rel.lower().endswith(".txt"):
            return os.path.join(root_dir, rel)

    return None


def md_to_docx_keep_images(md_file: str, out_docx: str):
    md_dir = os.path.dirname(md_file)
    input_format = "md" if md_file.lower().endswith(".md") else "markdown"
    pypandoc.convert_file(
        source_file=md_file,
        to="docx",
        format=input_format,
        outputfile=out_docx,
        extra_args=[
            "--standalone",
            f"--resource-path={md_dir}",
        ],
    )


def insert_translation_same_paragraph(docx_path: str, out_docx_translated: str, placeholder: str = "译文代替"):
    doc = Document(docx_path)
    for para in doc.paragraphs:
        if para.text.strip():
            para.add_run().add_break(WD_BREAK.LINE)
            para.add_run(placeholder)
    doc.save(out_docx_translated)


def process_pdf_to_bilingual_docx_core(pdf_path: str, pdf_output_dir: str) -> tuple[dict, int]:
    """
    核心处理逻辑（不依赖 Flask）：
      - 调 MinerUAPI，否则使用 LocalPDFProcessor 兜底
      - 下载/复制 zip，解压，查找 md/txt
      - md/txt -> docx（保图片）
      - 同段落插入“译文代替”
      - 返回 (payload_dict, status_code)
    """
    pdf_path = os.path.abspath(pdf_path)
    if not os.path.exists(pdf_path):
        return {"success": False, "error": f"未找到PDF: {pdf_path}"}, 400

    safe_mkdir(pdf_output_dir)

    # 1) 调 MinerU
    result = None
    try:
        from app.function.image_ocr.ocr_api import MinerUAPI
        logger.info("初始化 MinerUAPI ...")
        mineru_api = MinerUAPI()
        logger.info("MinerUAPI 初始化成功")
        result = mineru_api.process_pdf(pdf_path)
        logger.info(f"MinerU 处理结果: {result}")
    except Exception as e:
        logger.warning(f"MinerUAPI 处理失败: {e}")
        result = None

    # 2) 兜底本地处理
    if not result:
        try:
            from app.function.local_pdf_processor import LocalPDFProcessor
            local_processor = LocalPDFProcessor()
            result = local_processor.process_pdf(pdf_path)
            logger.info(f"本地PDF处理结果: {result}")
        except Exception as e:
            logger.error(f"本地处理器失败: {e}")
            logger.error(traceback.format_exc())
            return {"success": False, "error": "PDF处理失败"}, 500

    if not result:
        return {"success": False, "error": "所有PDF处理方法都失败了"}, 500

    if "code" in result and result["code"] != 0:
        return {"success": False, "error": f"PDF处理失败: {result.get('msg', '未知错误')}"}, 500

    if "data" not in result or "task_id" not in result["data"] or "full_zip_url" not in result["data"]:
        return {"success": False, "error": "返回数据缺少 task_id 或 full_zip_url"}, 500

    task_id = result["data"]["task_id"]
    zip_url = result["data"]["full_zip_url"]
    zip_path = os.path.join(pdf_output_dir, f"mineru_result_{task_id}.zip")

    try:
        download_or_copy_zip(zip_url, zip_path)
    except Exception as e:
        logger.error(f"获取ZIP失败: {e}")
        return {"success": False, "error": "获取ZIP失败"}, 500

    try:
        extract_zip(zip_path, pdf_output_dir)
    except Exception as e:
        logger.error(f"解压失败: {e}")
        return {"success": False, "error": "解压失败"}, 500

    md_file = find_md_or_txt(pdf_output_dir, task_id)
    if not md_file:
        return {"success": False, "error": "未找到MD/TXT"}, 500

    md_stem = os.path.splitext(os.path.basename(md_file))[0]
    docx_path = os.path.join(pdf_output_dir, f"{md_stem}.docx")
    bilingual_docx = os.path.join(pdf_output_dir, f"{md_stem}_translated.docx")

    try:
        md_to_docx_keep_images(md_file, docx_path)
        insert_translation_same_paragraph(docx_path, bilingual_docx)
    except Exception as e:
        logger.error(f"转换或插入译文失败: {e}")
        return {"success": False, "error": "转换失败"}, 500

    payload = {
        "success": True,
        "task_id": task_id,
        "md_file": os.path.relpath(md_file, pdf_output_dir),
        "docx_file": os.path.relpath(docx_path, pdf_output_dir),
        "bilingual_docx_file": os.path.relpath(bilingual_docx, pdf_output_dir),
        "output_dir": os.path.abspath(pdf_output_dir),
    }
    return payload, 0


# ========== 这里设置要处理的 PDF 路径和输出目录 ==========
if __name__ == "__main__":
    pdf_path = r"C:\Users\48846\Documents\123.pdf"  # 修改为你的PDF路径
    pdf_output_dir = r"D:\project\FCIAI2.0\app\uploads\pdf_outputs"  # 修改为你的输出目录

    payload, code = process_pdf_to_bilingual_docx_core(pdf_path, pdf_output_dir)
    import json
    print(json.dumps(payload, ensure_ascii=False, indent=2))
    sys.exit(code)
