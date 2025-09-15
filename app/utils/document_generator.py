#!/usr/bin/env python3
"""
文档生成器模块
用于创建双语Word文档，包含原文和译文的对照显示
"""

import os
import logging
from typing import List, Dict, Tuple
from docx import Document
from docx.shared import RGBColor, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import qn
from docx.enum.style import WD_STYLE_TYPE
import asyncio

# 为了按段落即时翻译，复用现有异步翻译能力
try:
    from app.function.local_qwen_async import translate_async
except Exception:
    translate_async = None

logger = logging.getLogger(__name__)


class BilingualDocumentGenerator:
    """
    双语文档生成器
    支持创建包含原文和译文对照的Word文档
    """
    
    def __init__(self):
        """初始化文档生成器"""
        self.document = Document()
        self._setup_styles()
    
    def _setup_styles(self):
        """设置文档样式"""
        # 设置默认字体
        style = self.document.styles['Normal']
        font = style.font
        font.name = '宋体'
        font._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        font.size = Pt(12)
        
        # 设置标题样式
        title_style = self.document.styles['Heading 1']
        title_font = title_style.font
        title_font.name = '黑体'
        title_font._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        title_font.size = Pt(16)
        title_font.bold = True
        
        # 设置二级标题样式
        heading2_style = self.document.styles['Heading 2']
        heading2_font = heading2_style.font
        heading2_font.name = '黑体'
        heading2_font._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        heading2_font.size = Pt(14)
        
    def add_heading(self, text: str, level: int = 1) -> None:
        """
        添加标题
        
        Args:
            text: 标题文本
            level: 标题级别 (1-6)
        """
        try:
            level = min(max(level, 1), 6)  # 限制在1-6之间
            self.document.add_heading(text, level=level)
            logger.info(f"添加标题: {text[:50]}...")
        except Exception as e:
            logger.error(f"添加标题失败: {e}")
            # 备用方法
            paragraph = self.document.add_paragraph()
            run = paragraph.add_run(text)
            run.bold = True
            run.font.size = Pt(16 - level)  # 级别越高字体越小
    
    def add_original_text(self, text: str) -> None:
        """
        添加原文段落
        
        Args:
            text: 原文文本
        """
        if not text.strip():
            self.document.add_paragraph()
            return
            
        try:
            paragraph = self.document.add_paragraph()
            run = paragraph.add_run(text.strip())
            # 原文使用黑色字体
            run.font.color.rgb = RGBColor(0, 0, 0)
            logger.info(f"添加原文: {text[:50]}...")
        except Exception as e:
            logger.error(f"添加原文失败: {e}")
            self.document.add_paragraph(text.strip())
    
    def add_translated_text(self, text: str) -> None:
        """
        添加译文段落
        
        Args:
            text: 译文文本
        """
        if not text.strip():
            return
            
        try:
            paragraph = self.document.add_paragraph()
            run = paragraph.add_run(text.strip())
            # 译文使用灰色字体以便区分，统一不使用斜体，保证风格一致
            run.font.color.rgb = RGBColor(96, 96, 96)
            run.italic = False
            
            # 设置段落格式
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(12)
            
            logger.info(f"添加译文: {text[:50]}...")
        except Exception as e:
            logger.error(f"添加译文失败: {e}")
            paragraph = self.document.add_paragraph()
            run = paragraph.add_run(text.strip())
            run.font.color.rgb = RGBColor(128, 128, 128)
    
    def add_bilingual_pair(self, original: str, translated: str) -> None:
        """
        添加原文和译文对
        
        Args:
            original: 原文
            translated: 译文
        """
        self.add_original_text(original)
        if translated and not translated.startswith('[翻译错误:') and translated.strip():
            self.add_translated_text(translated)
        self.document.add_paragraph()  # 添加空行分隔
    
    def add_list_item(self, text: str, numbered: bool = False) -> None:
        """
        添加列表项
        
        Args:
            text: 列表项文本
            numbered: 是否为有序列表
        """
        try:
            if numbered:
                self.document.add_paragraph(text.strip(), style='ListNumber')
            else:
                self.document.add_paragraph(text.strip(), style='ListBullet')
        except Exception as e:
            logger.warning(f"添加列表项失败，使用普通段落: {e}")
            self.add_original_text(text)
    
    def add_image(self, image_path: str, width_inches: float = 6.0) -> None:
        """
        添加图片到文档
        
        Args:
            image_path: 图片路径
            width_inches: 图片宽度（英寸）
        """
        try:
            if os.path.exists(image_path):
                from docx.shared import Inches
                self.document.add_picture(image_path, width=Inches(width_inches))
                logger.info(f"添加图片: {image_path}")
            else:
                logger.warning(f"图片文件不存在: {image_path}")
        except Exception as e:
            logger.error(f"添加图片失败: {e}")
    
    def save(self, file_path: str) -> bool:
        """
        保存文档
        
        Args:
            file_path: 保存路径
            
        Returns:
            bool: 是否保存成功
        """
        try:
            # 确保目录存在
            os.makedirs(os.path.dirname(file_path), exist_ok=True)
            
            # 保存文档
            self.document.save(file_path)
            logger.info(f"文档保存成功: {file_path}")
            return True
        except Exception as e:
            logger.error(f"保存文档失败: {e}")
            return False


def create_bilingual_word_document(
    content_lines: List[str], 
    translation_dict: Dict[str, str],
    output_path: str
) -> bool:
    """
    创建双语Word文档
    
    Args:
        content_lines: 原始内容行列表
        translation_dict: 翻译字典 {原文: 译文}
        output_path: 输出文件路径
        
    Returns:
        bool: 是否创建成功
    """
    try:
        generator = BilingualDocumentGenerator()
        
        i = 0
        while i < len(content_lines):
            line = content_lines[i].strip()
            
            if not line:
                generator.document.add_paragraph()
                i += 1
                continue
            
            # 处理标题 (# 开头)
            if line.startswith('#'):
                # 计算标题级别
                level = 1
                for char in line:
                    if char == '#':
                        level += 1
                    else:
                        break
                level = min(level, 6)  # 最多支持6级标题
                
                # 提取标题文本
                title_text = line.lstrip('# ').strip()
                generator.add_heading(title_text, level)
                
                # 检查下一行是否为译文
                if i + 1 < len(content_lines):
                    next_line = content_lines[i + 1].strip()
                    if next_line.startswith('【译文】'):
                        translated_text = next_line[5:].strip()  # 去掉"【译文】"前缀
                        generator.add_translated_text(translated_text)
                        i += 2  # 跳过标题和译文行
                        continue
                
                i += 1
                continue
            
            # 处理无序列表 (* 或 - 开头)
            if line.startswith('* ') or line.startswith('- '):
                list_text = line[2:].strip()
                generator.add_list_item(list_text, numbered=False)
                
                # 检查下一行是否为译文
                if i + 1 < len(content_lines):
                    next_line = content_lines[i + 1].strip()
                    if next_line.startswith('【译文】'):
                        translated_text = next_line[5:].strip()
                        generator.add_translated_text(translated_text)
                        i += 2  # 跳过列表项和译文行
                        continue
                
                i += 1
                continue
            
            # 处理有序列表 (数字. 开头)
            if line.startswith(tuple(f"{n}. " for n in range(1, 100))):  # 支持1-99的编号
                dot_index = line.find('. ')
                if dot_index > 0:
                    list_text = line[dot_index + 2:].strip()
                    generator.add_list_item(list_text, numbered=True)
                    
                    # 检查下一行是否为译文
                    if i + 1 < len(content_lines):
                        next_line = content_lines[i + 1].strip()
                        if next_line.startswith('【译文】'):
                            translated_text = next_line[5:].strip()
                            generator.add_translated_text(translated_text)
                            i += 2  # 跳过列表项和译文行
                            continue
                    
                    i += 1
                    continue
            
            # 处理译文标记行
            if line.startswith('【译文】'):
                # 这种情况理论上不应该出现，因为我们会在添加原文时处理对应的译文
                translated_text = line[5:].strip()
                generator.add_translated_text(translated_text)
                i += 1
                continue
            
            # 处理普通段落
            original_text = line
            translated_text = ""
            
            # 尝试查找翻译
            if original_text in translation_dict:
                translated_text = translation_dict[original_text]
            else:
                # 使用我们之前实现的模糊匹配逻辑
                for key, value in translation_dict.items():
                    if (original_text.strip() in key.strip() or 
                        key.strip() in original_text.strip()):
                        translated_text = value
                        break
            
            generator.add_bilingual_pair(original_text, translated_text)
            i += 1
        
        # 保存文档
        return generator.save(output_path)
        
    except Exception as e:
        logger.error(f"创建双语Word文档失败: {e}")
        import traceback
        logger.error(f"错误详情: {traceback.format_exc()}")
        return False


def process_markdown_to_bilingual_doc(
    markdown_content: str,
    translation_dict: Dict[str, str],
    output_path: str
) -> bool:
    """
    将Markdown内容处理为双语Word文档
    
    Args:
        markdown_content: Markdown格式的内容
        translation_dict: 翻译字典 {原文: 译文}
        output_path: 输出文件路径
        
    Returns:
        bool: 是否处理成功
    """
    try:
        # 按行分割内容
        content_lines = markdown_content.split('\n')
        
        # 创建双语文档
        success = create_bilingual_word_document(content_lines, translation_dict, output_path)
        
        if success:
            logger.info(f"成功创建双语Word文档: {output_path}")
        else:
            logger.error("创建双语Word文档失败")
            
        return success
        
    except Exception as e:
        logger.error(f"处理Markdown到双语文档失败: {e}")
        import traceback
        logger.error(f"错误详情: {traceback.format_exc()}")
        return False


def _sync_translate_single_text(text: str,
                                source_language: str = "en",
                                target_language: str = "zh") -> str:
    """
    使用现有的异步翻译接口对单条文本进行同步翻译，返回译文。

    若翻译不可用或失败，返回空字符串以便调用方降级处理。
    """
    if not text or not text.strip():
        return ""
    if translate_async is None:
        return ""
    try:
        async def _run_once() -> str:
            mapping = await translate_async(
                text.strip(),
                field=None,
                stop_words=[],
                custom_translations={},
                source_language=source_language,
                target_language=target_language
            )
            # translate_async 返回 {原文: 译文}
            if isinstance(mapping, dict):
                # 优先精确匹配整行
                if text.strip() in mapping:
                    return mapping[text.strip()] or ""
                # 退而求其次取第一个值
                for _, v in mapping.items():
                    return v or ""
            return ""

        # 独立运行事件循环
        try:
            return asyncio.run(_run_once())
        except RuntimeError:
            # 若已有事件循环（少见于Flask），则创建新循环执行
            loop = asyncio.new_event_loop()
            try:
                return loop.run_until_complete(_run_once())
            finally:
                loop.close()
    except Exception:
        return ""


def translate_markdown_to_bilingual_doc(
    markdown_content: str,
    output_path: str,
    source_language: str = "en",
    target_language: str = "zh",
    image_base_dir: str | None = None
) -> bool:
    """
    将Markdown内容按“标题/段落 → 逐条翻译 → 立即写入Word（原文在前，译文在后）”的方式生成双语Word文档。

    - 标题(`#`级别)写入为对应Heading，然后紧跟译文段落
    - 无序/有序列表项写入列表项，然后紧跟译文段落
    - 普通段落使用 add_bilingual_pair，原文后紧跟译文
    - 空行保持

    若翻译不可用或失败，依然写入原文，译文留空。
    """
    try:
        generator = BilingualDocumentGenerator()
        if not markdown_content:
            return generator.save(output_path)

        lines = markdown_content.split('\n')

        # 简单的中文检测（用于跳过已是中文的文本）
        def is_mostly_chinese(s: str) -> bool:
            if not s:
                return False
            total = len(s)
            if total == 0:
                return False
            import re
            cjk = len(re.findall(r"[\u4e00-\u9fff]", s))
            return (cjk / total) > 0.5

        # 解析 markdown 为块：标题、列表项、图片、普通段落（合并连续行）
        blocks: list[dict] = []
        i = 0
        while i < len(lines):
            raw_line = lines[i]
            line = (raw_line or "").rstrip('\r')

            # 跳过带有固定译文标记的行
            if line.strip().startswith('【译文】'):
                i += 1
                continue

            # 图片行（只处理整行图片语法）: ![alt](path)
            stripped = line.lstrip()
            if stripped.startswith('![') and '](' in stripped and stripped.endswith(')'):
                try:
                    alt_end = stripped.find('](')
                    path = stripped[alt_end+2:-1].strip()
                    blocks.append({'type': 'image', 'path': path, 'alt': stripped[2:alt_end]})
                    i += 1
                    continue
                except Exception:
                    pass

            # 标题
            if stripped.startswith('#'):
                level = 0
                for ch in stripped:
                    if ch == '#':
                        level += 1
                    else:
                        break
                level = max(1, min(level, 6))
                title_text = stripped[level:].strip()
                blocks.append({'type': 'heading', 'level': level, 'text': title_text})
                i += 1
                continue

            # 列表项
            if stripped.startswith('* ') or stripped.startswith('- '):
                item_text = stripped[2:].strip()
                blocks.append({'type': 'ul_item', 'text': item_text})
                i += 1
                continue

            dot_index = stripped.find('. ')
            if dot_index > 0 and stripped[:dot_index].isdigit():
                item_text = stripped[dot_index+2:].strip()
                blocks.append({'type': 'ol_item', 'text': item_text})
                i += 1
                continue

            # 空行 -> 作为段落分隔
            if not line.strip():
                blocks.append({'type': 'blank'})
                i += 1
                continue

            # 合并连续的非空普通文本行为一个段落
            paragraph_lines = [line.strip()]
            j = i + 1
            while j < len(lines):
                nxt = (lines[j] or "").rstrip('\r')
                nxt_stripped = nxt.lstrip()
                if not nxt.strip():
                    break
                if nxt_stripped.startswith('#') or nxt_stripped.startswith('![') or \
                   nxt_stripped.startswith('* ') or nxt_stripped.startswith('- '):
                    break
                di = nxt_stripped.find('. ')
                if di > 0 and nxt_stripped[:di].isdigit():
                    break
                if nxt_stripped.startswith('【译文】'):
                    break
                paragraph_lines.append(nxt.strip())
                j += 1
            merged = ' '.join(paragraph_lines).strip()
            if merged:
                blocks.append({'type': 'paragraph', 'text': merged})
            i = j if j > i else (i + 1)

        # 写入到 Word，使用去重缓存避免重复
        cache: dict[str, str] = {}
        for blk in blocks:
            btype = blk.get('type')
            if btype == 'blank':
                generator.document.add_paragraph()
                continue

            if btype == 'image':
                path = blk.get('path') or ''
                if path:
                    # 如果是相对路径，则基于 image_base_dir 解析
                    final_path = path
                    if image_base_dir and not (path.startswith('http://') or path.startswith('https://') or 
                                               path.startswith('/') or (len(path) > 1 and path[1] == ':')):
                        import os as _os
                        final_path = _os.path.join(image_base_dir, path)
                    try:
                        generator.document.add_picture(final_path, width=Inches(6.0))
                    except Exception:
                        # 图片失败时忽略，不阻断
                        pass
                continue

            text = blk.get('text', '').strip()
            if not text:
                continue

            # 跳过已是中文的原文，避免重复译文
            if is_mostly_chinese(text):
                generator.add_original_text(text)
                continue

            if btype == 'heading':
                level = int(blk.get('level', 1))
                generator.add_heading(text, level)
                translated = cache.get(text)
                if translated is None:
                    translated = _sync_translate_single_text(text, source_language, target_language)
                    cache[text] = translated
                if translated:
                    generator.add_translated_text(translated)
                generator.document.add_paragraph()
                continue

            if btype == 'ul_item':
                generator.add_list_item(text, numbered=False)
                translated = cache.get(text)
                if translated is None:
                    translated = _sync_translate_single_text(text, source_language, target_language)
                    cache[text] = translated
                if translated:
                    generator.add_translated_text(translated)
                generator.document.add_paragraph()
                continue

            if btype == 'ol_item':
                generator.add_list_item(text, numbered=True)
                translated = cache.get(text)
                if translated is None:
                    translated = _sync_translate_single_text(text, source_language, target_language)
                    cache[text] = translated
                if translated:
                    generator.add_translated_text(translated)
                generator.document.add_paragraph()
                continue

            # 普通段落
            translated = cache.get(text)
            if translated is None:
                translated = _sync_translate_single_text(text, source_language, target_language)
                cache[text] = translated
            generator.add_bilingual_pair(text, translated)

        return generator.save(output_path)
    except Exception as e:
        logger.error(f"按段落翻译并生成双语文档失败: {e}")
        import traceback
        logger.error(f"错误详情: {traceback.format_exc()}")
        return False